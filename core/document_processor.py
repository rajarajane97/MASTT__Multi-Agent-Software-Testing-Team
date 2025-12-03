"""
Document Processor for handling PDFs, Word docs, Confluence pages, etc.
"""

import os
import json
from pathlib import Path
from typing import List, Dict, Optional
import PyPDF2
import pdfplumber
from docx import Document
import markdown
from bs4 import BeautifulSoup
import requests
from atlassian import Confluence
from loguru import logger


class DocumentProcessor:
    """Processes various document formats for RAG ingestion."""
    
    def __init__(self, confluence_url: Optional[str] = None,
                 confluence_username: Optional[str] = None,
                 confluence_token: Optional[str] = None):
        """
        Initialize document processor.
        
        Args:
            confluence_url: Confluence instance URL
            confluence_username: Confluence username
            confluence_token: Confluence API token
        """
        self.processed_documents = []
        
        # Initialize Confluence client if credentials provided
        self.confluence = None
        if all([confluence_url, confluence_username, confluence_token]):
            try:
                self.confluence = Confluence(
                    url=confluence_url,
                    username=confluence_username,
                    password=confluence_token
                )
                logger.info("Confluence client initialized")
            except Exception as e:
                logger.warning(f"Could not initialize Confluence client: {str(e)}")
    
    def process_pdf(self, file_path: str) -> Dict:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dictionary with extracted content
        """
        try:
            content = {
                'source': file_path,
                'type': 'pdf',
                'text': '',
                'metadata': {}
            }
            
            # Try with pdfplumber first (better for complex PDFs)
            with pdfplumber.open(file_path) as pdf:
                content['metadata']['pages'] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        content['text'] += f"\n--- Page {page_num} ---\n{text}"
            
            logger.info(f"Processed PDF: {file_path}")
            return content
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    content['metadata']['pages'] = len(reader.pages)
                    
                    for page_num, page in enumerate(reader.pages, 1):
                        text = page.extract_text()
                        content['text'] += f"\n--- Page {page_num} ---\n{text}"
                
                return content
            except Exception as e2:
                logger.error(f"Fallback PDF processing also failed: {str(e2)}")
                return None
    
    def process_word_doc(self, file_path: str) -> Dict:
        """
        Extract text from Word document.
        
        Args:
            file_path: Path to Word file
            
        Returns:
            Dictionary with extracted content
        """
        try:
            doc = Document(file_path)
            
            content = {
                'source': file_path,
                'type': 'docx',
                'text': '',
                'metadata': {
                    'paragraphs': len(doc.paragraphs)
                }
            }
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content['text'] += para.text + '\n'
            
            # Extract tables
            if doc.tables:
                content['text'] += "\n--- Tables ---\n"
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join([cell.text for cell in row.cells])
                        content['text'] += row_text + '\n'
            
            logger.info(f"Processed Word document: {file_path}")
            return content
            
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {str(e)}")
            return None
    
    def process_markdown(self, file_path: str) -> Dict:
        """
        Extract text from Markdown file.
        
        Args:
            file_path: Path to Markdown file
            
        Returns:
            Dictionary with extracted content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # Convert to HTML then extract text
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text()
            
            content = {
                'source': file_path,
                'type': 'markdown',
                'text': text,
                'metadata': {
                    'raw_markdown': md_content[:500]  # Store snippet
                }
            }
            
            logger.info(f"Processed Markdown: {file_path}")
            return content
            
        except Exception as e:
            logger.error(f"Error processing Markdown {file_path}: {str(e)}")
            return None
    
    def process_text_file(self, file_path: str) -> Dict:
        """
        Extract text from plain text file.
        
        Args:
            file_path: Path to text file
            
        Returns:
            Dictionary with extracted content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            content = {
                'source': file_path,
                'type': 'text',
                'text': text,
                'metadata': {
                    'lines': len(text.splitlines())
                }
            }
            
            logger.info(f"Processed text file: {file_path}")
            return content
            
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {str(e)}")
            return None
    
    def fetch_confluence_page(self, page_id: str) -> Dict:
        """
        Fetch content from Confluence page.
        
        Args:
            page_id: Confluence page ID
            
        Returns:
            Dictionary with extracted content
        """
        if not self.confluence:
            logger.error("Confluence client not initialized")
            return None
        
        try:
            page = self.confluence.get_page_by_id(
                page_id,
                expand='body.storage,version'
            )
            
            # Extract HTML content
            html_content = page['body']['storage']['value']
            soup = BeautifulSoup(html_content, 'html.parser')
            text = soup.get_text()
            
            content = {
                'source': f"confluence:{page_id}",
                'type': 'confluence',
                'text': text,
                'metadata': {
                    'title': page['title'],
                    'version': page['version']['number'],
                    'page_id': page_id
                }
            }
            
            logger.info(f"Fetched Confluence page: {page['title']}")
            return content
            
        except Exception as e:
            logger.error(f"Error fetching Confluence page {page_id}: {str(e)}")
            return None
    
    def fetch_confluence_space(self, space_key: str, limit: int = 50) -> List[Dict]:
        """
        Fetch all pages from a Confluence space.
        
        Args:
            space_key: Confluence space key
            limit: Maximum number of pages to fetch
            
        Returns:
            List of documents
        """
        if not self.confluence:
            logger.error("Confluence client not initialized")
            return []
        
        documents = []
        
        try:
            pages = self.confluence.get_all_pages_from_space(
                space_key,
                start=0,
                limit=limit,
                expand='body.storage'
            )
            
            for page in pages:
                doc = self.fetch_confluence_page(page['id'])
                if doc:
                    documents.append(doc)
            
            logger.info(f"Fetched {len(documents)} pages from space {space_key}")
            return documents
            
        except Exception as e:
            logger.error(f"Error fetching Confluence space {space_key}: {str(e)}")
            return []
    
    def process_directory(self, directory_path: str) -> List[Dict]:
        """
        Process all supported documents in a directory.
        
        Args:
            directory_path: Path to directory
            
        Returns:
            List of processed documents
        """
        documents = []
        directory = Path(directory_path)
        
        if not directory.exists():
            logger.error(f"Directory not found: {directory_path}")
            return documents
        
        # Supported file extensions
        processors = {
            '.pdf': self.process_pdf,
            '.docx': self.process_word_doc,
            '.doc': self.process_word_doc,
            '.md': self.process_markdown,
            '.txt': self.process_text_file,
            '.rst': self.process_text_file
        }
        
        for file_path in directory.rglob('*'):
            if file_path.is_file():
                ext = file_path.suffix.lower()
                
                if ext in processors:
                    doc = processors[ext](str(file_path))
                    if doc:
                        documents.append(doc)
        
        logger.info(f"Processed {len(documents)} documents from {directory_path}")
        return documents
    
    def process_file(self, file_path: str) -> Optional[Dict]:
        """
        Process a single file.
        
        Args:
            file_path: Path to file
            
        Returns:
            Processed document or None
        """
        file_path = Path(file_path)
        ext = file_path.suffix.lower()
        
        processors = {
            '.pdf': self.process_pdf,
            '.docx': self.process_word_doc,
            '.doc': self.process_word_doc,
            '.md': self.process_markdown,
            '.txt': self.process_text_file,
            '.rst': self.process_text_file
        }
        
        if ext in processors:
            return processors[ext](str(file_path))
        else:
            logger.warning(f"Unsupported file type: {ext}")
            return None
    
    def save_processed_documents(self, output_path: str):
        """Save processed documents to JSON file."""
        output_file = Path(output_path) / 'processed_documents.json'
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_file, 'w') as f:
            json.dump(self.processed_documents, f, indent=2)
        
        logger.info(f"Processed documents saved to {output_file}")
